import os
import sys
from pathlib import Path

# Add core to path
sys.path.append(str(Path(__file__).parent.parent))

from core.document_converter import convert_docx_to_pdf
from docx import Document

def test_pipeline():
    test_docx = Path("test_sample.docx")
    test_pdf = Path("test_sample.pdf")
    
    # 1. Create a dummy docx
    print("Creating dummy DOCX...")
    doc = Document()
    doc.add_heading('References', 0)
    doc.add_paragraph('1. Doe, J. (2020). Test Article. Journal of Testing, 1(1), 1-10. doi:10.1001/test.123')
    doc.add_paragraph('2. Smith, A. (2021). Another Test. Testing Journal, 2(2), 20-30.')
    doc.save(test_docx)
    
    # 2. Test conversion
    print(f"Testing conversion: {test_docx} -> {test_pdf}")
    success = convert_docx_to_pdf(str(test_docx), str(test_pdf))
    
    if success:
        print("SUCCESS: Conversion completed.")
        if test_pdf.exists():
            print(f"PDF file exists: {test_pdf} ({test_pdf.stat().st_size} bytes)")
            test_pdf.unlink()
    else:
        print("FAILED: Conversion failed (Note: This is expected if Word is not installed/running in this environment).")
    
    if test_docx.exists():
        test_docx.unlink()

if __name__ == "__main__":
    test_pipeline()
